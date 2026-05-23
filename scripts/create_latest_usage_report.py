from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/StudyPulse_latest_usage_report.docx")


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, "F4EFE6")
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), "E8D9C7")
        borders.append(el)
    tc_pr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color="1F4F46")
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_key_value_table(doc):
    rows = [
        ("软件名称", "StudyPulse"),
        ("当前版本", "0.20.0"),
        ("适用平台", "Windows 桌面端"),
        ("安装包", "StudyPulse_0.20.0_x64_安装包.zip"),
        ("主要技术栈", "Tauri 2、React、TypeScript、Vite、TailwindCSS、SQLite、Recharts、Rust Windows API"),
        ("数据保存方式", "默认本地 SQLite 保存"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["项目", "内容"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        shade_cell(cell, "DDEBE5")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=9.5, bold=True)
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    set_run_font(run, size=9.5)
    for row in table.rows:
        set_cell_width(row.cells[0], 1800)
        set_cell_width(row.cells[1], 7200)
    doc.add_paragraph()


def add_validation_table(doc):
    rows = [
        ("cargo check", "通过", "Rust/Tauri 后端基础编译检查通过。"),
        ("cargo test", "通过", "Rust 单元测试通过，覆盖数据库、番茄钟、AI 错误处理、活跃度等逻辑。"),
        ("npm test", "通过", "前端基础渲染测试通过。"),
        ("npm run build", "通过", "TypeScript 编译和 Vite 生产构建通过。"),
        ("MSI 打包", "通过", "已生成 0.20.0 MSI 安装包。"),
        ("签名验证", "通过", "MSI 已使用 StudyPulse Test Code Signing 自签名证书签名，并通过 signtool verify。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["验证项", "结果", "说明"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        shade_cell(cell, "DDEBE5")
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=9, bold=True)
    for item, result, note in rows:
        cells = table.add_row().cells
        values = [item, result, note]
        for i, value in enumerate(values):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    set_run_font(run, size=9)
    for row in table.rows:
        set_cell_width(row.cells[0], 1900)
        set_cell_width(row.cells[1], 1000)
        set_cell_width(row.cells[2], 6100)


def build_document():
    OUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.line_spacing = 1.25
    styles["Normal"].paragraph_format.space_after = Pt(5)

    for style_name, size, color in [
        ("Heading 1", 16, "1F4F46"),
        ("Heading 2", 13, "2F6F5E"),
        ("Heading 3", 11.5, "222222"),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    muted = RGBColor(90, 95, 91)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("StudyPulse 最新版使用报告")
    set_run_font(r, size=22, bold=True, color="1F4F46")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("版本 0.20.0 | Windows 桌面端学习辅助工具")
    set_run_font(r, size=11)
    r.font.color.rgb = muted

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("文档日期：2026年5月23日   |   文档类型：使用报告")
    set_run_font(r, size=10)
    r.font.color.rgb = muted

    add_callout(
        doc,
        "报告结论：StudyPulse 0.20.0 已具备可演示和可安装的 MVP 形态，能够完成学习会话记录、番茄钟、窗口采集、活跃度统计、本地日报、历史日报和 AI 总结等核心流程。",
    )

    doc.add_heading("一、基本信息", level=1)
    add_key_value_table(doc)

    doc.add_heading("二、项目简介", level=1)
    doc.add_paragraph(
        "StudyPulse 是一个轻量级 Windows 桌面学习辅助工具。用户打开程序并点击开始学习后，程序会记录学习期间的会话状态、当前前台窗口、应用使用时长、键鼠活跃度和番茄钟状态。学习结束后，程序会生成本地日报，并可根据用户配置的 OpenAI 兼容 API 生成学习总结。"
    )
    doc.add_paragraph(
        "该项目不做账号系统和云同步，重点是本地可运行、界面简洁、采集边界清晰，适合作为课程项目、学习工具原型或桌面端 MVP 展示。"
    )

    doc.add_heading("三、最新版更新内容", level=1)
    add_bullets(
        doc,
        [
            "版本号统一升级为 0.20.0，安装包文件名同步更新。",
            "新增历史日报入口，可以查看最近学习记录。",
            "首次启动增加隐私说明，帮助用户理解采集边界。",
            "番茄钟支持 25、40、50 分钟预设和自定义时长。",
            "AI 总结支持温和鼓励、正常复盘、轻微吐槽、严格监督四种语气。",
            "替换了正式一些的应用图标，并重新生成 MSI 安装包。",
        ],
    )

    doc.add_heading("四、安装与启动方式", level=1)
    doc.add_paragraph("当前推荐使用安装包方式启动，适合交给其他电脑安装测试。")
    add_bullets(
        doc,
        [
            "解压 StudyPulse_0.20.0_x64_安装包.zip。",
            "可选：先导入 StudyPulse-Test-Code-Signing.cer 到当前用户的受信任根证书颁发机构。",
            "双击 StudyPulse_Setup_0.20.0_x64.msi 完成安装。",
            "安装完成后，从 Windows 开始菜单搜索 StudyPulse 并打开。",
            "开发调试时，也可以在项目根目录运行 npm run tauri dev。",
        ],
    )

    doc.add_heading("五、主要功能使用说明", level=1)
    add_bullets(
        doc,
        [
            "学习会话：点击开始学习后进入记录状态，点击结束学习后停止采集并生成日报。",
            "番茄钟：支持预设和自定义时长，可以开始、暂停、继续和重置。",
            "前台窗口采集：Windows 下每秒采样当前前台应用、窗口标题和可执行路径。",
            "应用排行：根据采样数据统计应用使用时长，帮助用户复盘主要时间花在哪里。",
            "键鼠活跃度：只统计键盘和鼠标事件数量，不保存具体按键、输入内容或鼠标坐标。",
            "本地日报与历史日报：学习结束后生成本地日报，历史日报入口可回看最近记录。",
            "AI 总结与聊天：可根据日报生成不同语气的总结，并继续追问学习状态。",
        ],
    )

    doc.add_heading("六、实际使用流程记录", level=1)
    doc.add_paragraph(
        "一次典型使用流程为：启动 StudyPulse，阅读首次隐私说明，点击开始学习；学习过程中切换文档、浏览器或代码编辑器，并使用番茄钟辅助计时；结束学习后点击结束学习，程序生成日报；用户可在历史日报中回看记录，也可以选择 AI 总结语气并生成复盘反馈。"
    )
    doc.add_paragraph(
        "从当前版本看，首页已经能覆盖今日学习时间、当前状态、当前活跃应用、专注度、应用排行、活跃度趋势和 AI 总结区，整体更接近一个可试用的学习仪表盘。"
    )

    doc.add_heading("七、功能验证结果", level=1)
    add_validation_table(doc)

    doc.add_heading("八、隐私与数据保存说明", level=1)
    add_bullets(
        doc,
        [
            "所有学习数据默认保存在本机 SQLite 数据库中。",
            "程序不记录具体按键，不记录输入内容，不记录鼠标坐标。",
            "程序不截图、不录屏，也不会主动上传本地采集数据。",
            "只有用户主动生成 AI 总结或发送聊天消息时，日报摘要和聊天内容才会发送到用户配置的 API。",
        ],
    )

    doc.add_heading("九、当前问题与限制", level=1)
    add_bullets(
        doc,
        [
            "当前主要支持 Windows，其他系统暂未实现同等采集能力。",
            "安装包使用自签名测试证书，不能完全消除 Windows SmartScreen 提示。",
            "本机 WiX 在 ICE 校验阶段存在 Windows Installer 服务问题，因此 MSI 通过手动 light.exe -sval 方式生成。",
            "AI 总结依赖用户配置或默认 API 服务状态，网络或服务异常时可能失败。",
            "历史日报已有入口，但更复杂的周报、月报、应用分类和目标管理仍属于后续功能。",
        ],
    )

    doc.add_heading("十、后续改进建议", level=1)
    add_bullets(
        doc,
        [
            "增加周报和月报视图，展示长期学习趋势。",
            "增加应用分类或白名单/黑名单，让专注度评分更符合真实学习场景。",
            "优化安装包签名，后续公开分发时考虑正式代码签名证书。",
            "进一步简化 AI 配置流程，给非技术用户更清楚的默认选项。",
            "补充更多前端交互测试和真实安装场景测试。",
        ],
    )

    doc.add_heading("十一、总体评价", level=1)
    doc.add_paragraph(
        "StudyPulse 0.20.0 已经完成从前端原型到 Windows 桌面 MVP 的基本闭环。它能完成学习开始、后台统计、结束日报、AI 总结和历史回看等关键流程，同时保持本地保存和隐私边界说明。当前版本适合课程展示、同学试用和继续迭代。"
    )
    doc.add_paragraph(
        "下一阶段的重点不应是扩大功能范围，而是提高数据解释能力和长期使用价值，例如应用分类、周报/月报、目标备注以及正式签名交付。"
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("StudyPulse 0.20.0 使用报告")
    set_run_font(r, size=9)
    r.font.color.rgb = muted

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build_document()
