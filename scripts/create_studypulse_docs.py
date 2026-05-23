from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs"
DATE_TEXT = "2026年5月23日"

INK = RGBColor(31, 41, 51)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MOSS = RGBColor(47, 111, 94)
MUTED = RGBColor(102, 114, 127)
LIGHT_FILL = "F2F4F7"


def set_east_asia_font(run, font_name: str = "Microsoft YaHei UI") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_east_asia_font(run)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    tail = paragraph.add_run(" 页")
    set_east_asia_font(tail)


def setup_document(title: str, doc_type: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title_style = styles["Title"]
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.font.color.rgb = INK
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
    title_style.paragraph_format.space_after = Pt(4)

    subtitle_style = styles["Subtitle"]
    subtitle_style.font.name = "Calibri"
    subtitle_style.font.size = Pt(11)
    subtitle_style.font.color.rgb = MUTED
    subtitle_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
    subtitle_style.paragraph_format.space_after = Pt(14)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer_run = footer.add_run("StudyPulse 项目材料    ")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = MUTED
    set_east_asia_font(footer_run)
    add_page_number(footer)

    p = doc.add_paragraph(style="Title")
    run = p.add_run(title)
    set_east_asia_font(run)
    p = doc.add_paragraph(style="Subtitle")
    run = p.add_run(f"{doc_type} | {DATE_TEXT}")
    set_east_asia_font(run)

    return doc


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_east_asia_font(run)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        run = p.add_run(item)
        set_east_asia_font(run)


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(6.5)
    set_cell_margins(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = MOSS
    set_east_asia_font(r)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_east_asia_font(r)


def add_meta_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(1.65)
    table.columns[1].width = Inches(4.85)
    set_cell_margins(table)
    hdr = table.rows[0].cells
    hdr[0].text = "项目"
    hdr[1].text = "内容"
    for cell in hdr:
        set_cell_shading(cell, LIGHT_FILL)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                set_east_asia_font(run)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for cell in cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_east_asia_font(run)


def project_intro() -> None:
    doc = setup_document("StudyPulse 项目简介", "项目文字描述")
    add_callout(
        doc,
        "一句话概述",
        "StudyPulse 是一个面向 Windows 桌面端的轻量级学习辅助工具，帮助用户记录学习会话中的软件使用、番茄钟状态、键鼠活跃度，并在结束后生成本地日报和 AI 复盘。",
    )

    doc.add_heading("一、项目背景", level=1)
    add_para(
        doc,
        "大学生在自习、写代码、看网课或查资料时，经常会在多个软件和网页之间切换。传统番茄钟只能记录时间，无法回答“这段时间主要花在什么软件上”“中途是否频繁切换”“学习结束后有什么值得复盘”等问题。StudyPulse 的设计目标是用尽量轻量的方式，把学习过程转化成可查看、可复盘的本地数据。",
    )

    doc.add_heading("二、核心目标", level=1)
    add_bullets(
        doc,
        [
            "打开程序并点击开始后，即进入一次学习会话，结束时生成本地日报。",
            "记录前台窗口和应用使用时长，帮助用户了解学习期间的注意力分布。",
            "统计键盘和鼠标活跃度，只保存数量，不记录具体输入内容。",
            "通过番茄钟和 AI 总结，让用户获得更容易理解的学习反馈。",
            "所有采集数据默认保存在本机，避免做成复杂的账号或云同步系统。",
        ],
    )

    doc.add_heading("三、主要功能", level=1)
    add_meta_table(
        doc,
        [
            ("学习会话", "支持开始、结束学习，并记录会话起止时间和状态。"),
            ("番茄钟", "支持开始、暂停、重置和完成统计，帮助用户维持学习节奏。"),
            ("窗口采集", "Windows 后端采集当前前台应用、窗口标题、可执行路径和采样时间。"),
            ("应用排行", "根据窗口采样数据聚合应用使用时长，生成常用软件排行。"),
            ("键鼠活跃度", "只统计键盘和鼠标事件数量，用于显示学习过程中的活跃趋势。"),
            ("本地日报", "结束学习后汇总学习时长、应用排行、活跃度、番茄钟完成数和专注度。"),
            ("AI 总结与聊天", "用户主动触发后，将日报摘要发送给配置的 OpenAI 兼容 API 生成复盘。"),
        ],
    )

    doc.add_heading("四、技术实现", level=1)
    add_para(
        doc,
        "项目采用 Tauri 2 + React + TypeScript + Vite + TailwindCSS 构建桌面应用界面，后端使用 Rust 实现本地数据处理、Windows 前台窗口采集、键鼠活跃度计数和 AI 请求。数据默认存储在 SQLite 中，图表展示使用 Recharts。",
    )
    add_bullets(
        doc,
        [
            "前端负责仪表盘、番茄钟控制、AI 总结区和设置弹窗。",
            "Rust 后端负责 Tauri commands、SQLite 初始化、会话管理和本地采集。",
            "Windows 采集只在学习会话进行中启用，结束后停止采样任务。",
            "AI 配置采用 base_url、api_key、model 的 OpenAI 兼容格式，查询时不返回明文 API key。",
        ],
    )

    doc.add_heading("五、隐私边界", level=1)
    add_callout(
        doc,
        "隐私说明",
        "StudyPulse 不记录具体按键内容，不记录输入文本，不截图，不记录鼠标坐标，也不会自动上传本地采集数据。只有当用户主动点击生成 AI 总结或发送聊天时，日报摘要才会发送给用户配置的 AI API。",
    )

    doc.add_heading("六、当前状态与后续方向", level=1)
    add_para(
        doc,
        "当前版本已经具备 MVP 的主要闭环：会话记录、番茄钟、窗口采集、应用排行、键鼠计数、本地日报、AI 总结和 Windows MSI 安装包。后续可以继续优化历史日报浏览、应用分类、AI 总结风格、正式图标、安装包签名和新手引导。",
    )
    add_bullets(
        doc,
        [
            "增加历史日报列表，便于按周或按月查看学习趋势。",
            "支持应用白名单、黑名单或学习/娱乐分类，提升专注度评分的解释性。",
            "允许用户选择 AI 总结语气，例如鼓励型、严肃型或轻微吐槽型。",
            "完善首次启动说明，让用户更清楚地理解采集范围和隐私边界。",
            "替换正式图标，并使用可信代码签名证书优化安装体验。",
        ],
    )

    doc.save(OUT / "StudyPulse_项目简介.docx")


def user_feedback_a() -> None:
    doc = setup_document("StudyPulse 用户体验反馈 - 用户A", "虚拟用户访谈记录")
    add_meta_table(
        doc,
        [
            ("用户身份", "大三计算机科学与技术专业学生"),
            ("主要场景", "写课程实验、刷算法题、看编程网课、整理项目报告"),
            ("使用设备", "Windows 笔记本电脑"),
            ("体验周期", "模拟体验 2 次，每次约 40-60 分钟"),
        ],
    )

    doc.add_heading("一、使用场景", level=1)
    add_para(
        doc,
        "我平时学习主要集中在 VS Code、浏览器、终端和文档软件之间。以前我只会用普通番茄钟记录自己坐了多久，但很难判断这段时间到底是在写代码、查资料，还是不知不觉切去网页刷东西。StudyPulse 对我来说比较有用的地方，是它能把这些应用使用情况汇总出来。",
    )

    doc.add_heading("二、实际体验", level=1)
    add_para(
        doc,
        "开始学习后，首页能直接看到今日学习时长、当前状态、当前窗口、应用排行和键鼠活跃度。界面没有太多复杂动画，比较像一个学习仪表盘。结束学习后生成日报，再让 AI 做一个简短总结，这一点比单纯看数字更容易理解。",
    )
    add_para(
        doc,
        "我比较喜欢应用使用排行，因为它能看出自己到底在 Code 里待了多久、浏览器占了多少时间。对于写代码的人来说，终端、编辑器和浏览器本来都会频繁切换，所以如果后面能支持应用分类，就能更准确地区分“查资料”和“分心”。",
    )

    doc.add_heading("三、满意点", level=1)
    add_bullets(
        doc,
        [
            "学习会话和番茄钟放在一起，使用逻辑比较直观。",
            "应用排行对编程学习很有帮助，可以看到时间主要花在哪些工具上。",
            "键鼠活跃度只统计数量，不记录具体按键，这一点让我比较放心。",
            "AI 总结的语气如果保持轻微吐槽和鼓励，会比普通日报更有阅读动力。",
        ],
    )

    doc.add_heading("四、遇到的问题", level=1)
    add_bullets(
        doc,
        [
            "第一次使用时不太确定窗口标题会记录到什么程度，希望隐私说明能更醒目。",
            "目前缺少历史日报入口，结束后如果没有及时看，之后不太方便回顾。",
            "应用排行只显示软件名称还不够，希望能区分学习类软件和娱乐类软件。",
            "安装包如果没有正式签名，给同学安装时可能会被 Windows 安全提示吓到。",
        ],
    )

    doc.add_heading("五、修改建议", level=1)
    add_bullets(
        doc,
        [
            "增加历史日报列表，至少可以查看最近 7 天的学习时长和应用排行。",
            "增加应用分类功能，例如把 VS Code、终端、PDF、浏览器资料页归为学习工具。",
            "增加白名单或黑名单，让用户自己定义哪些软件会降低专注度评分。",
            "番茄钟时长最好可以自定义，比如 25 分钟、40 分钟、50 分钟几种常用模式。",
            "AI 总结可以提供语气选项，例如“严格复盘”“正常鼓励”“轻微吐槽”。",
        ],
    )

    doc.add_heading("六、总体评价", level=1)
    add_para(
        doc,
        "整体来看，StudyPulse 对我这种经常在电脑上学习和写代码的学生是有价值的。它不需要我手动记录太多东西，但学习结束后能给出一个大致复盘。当前版本已经能跑通主要流程，后续如果补上历史记录和应用分类，我会更愿意长期使用。",
    )

    doc.save(OUT / "StudyPulse_用户体验反馈_用户A.docx")


def user_feedback_b() -> None:
    doc = setup_document("StudyPulse 用户体验反馈 - 用户B", "虚拟用户访谈记录")
    add_meta_table(
        doc,
        [
            ("用户身份", "大二英语专业学生"),
            ("主要场景", "背单词、阅读英文资料、写课程论文、整理课堂笔记"),
            ("使用设备", "Windows 轻薄本"),
            ("体验周期", "模拟体验 3 次，每次约 30 分钟"),
        ],
    )

    doc.add_heading("一、使用场景", level=1)
    add_para(
        doc,
        "我平时学习不是写代码，更多是打开 Word、PDF 阅读器、浏览器和背单词软件。对我来说，学习时最难的是发现自己什么时候分心了，因为有时候只是查一个词，最后就会在网页上停很久。StudyPulse 的应用排行能让我看到自己是不是把太多时间花在浏览器上。",
    )

    doc.add_heading("二、实际体验", level=1)
    add_para(
        doc,
        "界面整体比较干净，打开后能看到今日学习时间和当前状态。番茄钟区域比较容易理解，我可以把它当成普通番茄钟来用。结束学习后生成 AI 总结这一点对我比较有吸引力，因为我不太想自己再写一段复盘。",
    )
    add_para(
        doc,
        "不过我一开始对“窗口采集”和“键鼠活跃度”有点紧张，会担心是不是在记录我打了什么字。后来看到它说明不记录具体按键、不记录输入内容、不截图，才放心一些。这个说明如果在第一次启动时就弹出来，会更好。",
    )

    doc.add_heading("三、满意点", level=1)
    add_bullets(
        doc,
        [
            "番茄钟和学习时间放在首页，使用起来不需要学习成本。",
            "日报可以帮我回忆这段时间主要在做什么，不用靠主观感觉判断。",
            "AI 总结比普通数据更亲切，适合学习结束后快速看一眼。",
            "数据默认保存在本地，对学生用户来说比较安心。",
        ],
    )

    doc.add_heading("四、使用中的困惑", level=1)
    add_bullets(
        doc,
        [
            "专注度分数是怎么算的还不够清楚，容易误以为它是很准确的评价。",
            "浏览器既可能是查资料，也可能是分心，希望软件不要简单地把浏览器都算成低效。",
            "设置页里的 API URL、模型名、API Key 对非技术专业学生来说比较陌生。",
            "AI 总结如果过于吐槽，可能会让人有压力，希望语气能自己选择。",
        ],
    )

    doc.add_heading("五、修改建议", level=1)
    add_bullets(
        doc,
        [
            "增加新手引导，用简单语言说明软件记录什么、不记录什么。",
            "把 AI 设置做得更友好，可以提供默认选项和“我不懂，先跳过”的按钮。",
            "增加学习目标备注，例如本次学习目标是“背 100 个单词”或“完成论文大纲”。",
            "日报里增加一句“今天最值得肯定的地方”，让反馈更鼓励一些。",
            "支持历史日报列表，方便我看一周内有没有坚持学习。",
        ],
    )

    doc.add_heading("六、总体评价", level=1)
    add_para(
        doc,
        "我觉得 StudyPulse 比普通番茄钟更适合电脑学习，因为它能把软件使用情况和学习时间结合起来。对非技术专业学生来说，最重要的是隐私说明和设置流程要更简单。如果能把这些地方做得更清楚，我会愿意把它当成日常自习工具使用。",
    )

    doc.save(OUT / "StudyPulse_用户体验反馈_用户B.docx")


def main() -> None:
    OUT.mkdir(exist_ok=True)
    project_intro()
    user_feedback_a()
    user_feedback_b()
    print("created:")
    for path in sorted(OUT.glob("StudyPulse_*.docx")):
        print(path)


if __name__ == "__main__":
    main()
