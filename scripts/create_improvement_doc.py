from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/StudyPulse_improvement_plan.docx")


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, "F4EFE6")

    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), "E8D9C7")
        borders.append(el)
    tc_pr.append(borders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color="1F4F46")
    doc.add_paragraph()


def build_document():
    OUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.line_spacing = 1.25
    styles["Normal"].paragraph_format.space_after = Pt(5)

    for style_name, size, color in [
        ("Heading 1", 16, "1F4F46"),
        ("Heading 2", 13, "2F6F5E"),
        ("Heading 3", 11.5, "222222"),
    ]:
        st = styles[style_name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(5)

    muted = RGBColor(90, 95, 91)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("StudyPulse 修改建议与改进方案")
    set_run_font(r, size=22, bold=True, color="1F4F46")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于两位大学生用户体验反馈的产品迭代说明")
    set_run_font(r, size=11)
    r.font.color.rgb = muted

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("文档日期：2026年5月23日   |   文档类型：改进方案")
    set_run_font(r, size=10)
    r.font.color.rgb = muted

    add_callout(
        doc,
        "核心结论：当前 StudyPulse 的 MVP 功能已经具备可用基础。下一步应优先解决长期复盘、新用户信任感、AI 体验和安装交付问题，而不是继续堆叠复杂功能。",
    )

    sections = [
        (
            "一、项目当前状态简述",
            [
                "StudyPulse 是一个 Windows 桌面学习辅助工具，当前版本已经支持学习会话、番茄钟、前台窗口采集、应用使用时长统计、键鼠活跃度计数、本地日报、AI 总结与聊天。",
                "数据默认保存在本地 SQLite 中。程序不记录具体按键、不记录输入内容、不截图，也不会主动上传本地采集数据。只有用户主动生成 AI 总结或继续聊天时，日报摘要才会发送到用户配置的 API。",
                "从两位用户反馈看，产品方向是成立的，但还需要增强复盘连续性、隐私解释、低门槛设置和安装可信度。",
            ],
        ),
        (
            "二、用户 A/B 反馈要点归纳",
            [
                "用户 A 是计算机专业学生，主要场景是写代码、看网课、刷算法题和整理实验报告。他更关注应用排行、代码工具与浏览器的区分、历史记录和番茄钟灵活性。",
                "用户 B 是英语专业学生，主要场景是背单词、阅读资料、写论文和整理笔记。她更关注隐私说明、设置门槛、AI 总结语气和日报是否足够鼓励。",
                "两位用户共同提到的重点包括：希望查看历史日报，希望隐私说明更醒目，希望 AI 总结可调语气，希望番茄钟更灵活，并希望安装包减少安全提示。",
            ],
        ),
    ]
    for title, paragraphs in sections:
        doc.add_heading(title, level=1)
        for text in paragraphs:
            doc.add_paragraph(text)

    doc.add_heading("三、保留的五项修改建议", level=1)
    items = [
        (
            "1. 增加历史日报列表",
            "最高",
            "用户结束学习后需要回看过往记录，否则日报只能当场使用，难以形成长期复盘。",
            "新增历史日报入口，默认展示最近 7 天或最近 30 天记录。每条日报包含学习时长、专注度、应用排行、番茄钟完成数和 AI 总结。数据从本地 SQLite 的 daily_reports 表读取，不引入账号或云同步。",
            "让 StudyPulse 从一次性计时工具变成可持续使用的学习复盘工具。",
        ),
        (
            "2. 增加首次启动隐私引导",
            "最高",
            "用户对窗口采集和键鼠活跃度容易产生担心，需要在第一次使用前建立信任。",
            "首次启动时弹出简短说明，明确记录当前前台应用和窗口标题，只统计键盘、鼠标事件数量，不记录具体按键、输入内容、鼠标坐标、截图或录屏。用户确认后不再反复弹出，并可在设置页重新查看。",
            "降低用户心理负担，提高非技术用户的接受度。",
        ),
        (
            "4. AI 总结增加语气选择",
            "高",
            "AI 总结是产品亮点，但不同用户对“吐槽”和“鼓励”的接受程度不同。",
            "提供“温和鼓励、正常复盘、轻微吐槽、严格监督”四种语气选项。生成总结时将语气写入 prompt。未配置 API 时继续返回本地 mock 总结，不影响日报使用。",
            "让 AI 反馈更贴合用户性格，避免过度吐槽造成压力。",
        ),
        (
            "5. 番茄钟支持自定义时长",
            "中",
            "不同专业和任务节奏不同，固定 25 分钟不能覆盖所有学习场景。",
            "提供 25、40、50 分钟预设，并允许用户输入自定义分钟数。保存上次使用的时长，日报继续统计番茄钟完成数量。",
            "提升番茄钟的适配性，让用户更愿意长期使用。",
        ),
        (
            "7. 优化安装包签名和图标",
            "中",
            "当前安装包使用自签名测试证书，可能出现 Windows 安全提示，影响他人试用。",
            "替换正式应用图标，优化 MSI 文件名，例如 StudyPulse_Setup_0.1.0_x64.msi。继续保留自签名证书测试方案，并在使用手册中说明自签名证书不能完全消除 SmartScreen；公开发布需要正式代码签名证书。",
            "提升安装可信度和项目完成度，便于给同学或老师演示。",
        ),
    ]

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["建议", "优先级", "改进原因", "具体方案", "预期效果"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, "DDEBE5")
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=9, bold=True)

    for row in items:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    set_run_font(run, size=9)

    widths = [1500, 900, 2100, 3300, 2100]
    for row in table.rows:
        for idx, width in enumerate(widths):
            set_cell_width(row.cells[idx], width)

    doc.add_heading("四、分阶段实施计划", level=1)
    phases = [
        ("第一阶段：提升可用性与信任感", ["实现历史日报列表。", "实现首次启动隐私引导。", "在设置页增加隐私说明入口。"]),
        ("第二阶段：优化 AI 与番茄钟体验", ["实现 AI 总结语气选择。", "调整 AI prompt，使不同语气有明显差异。", "实现番茄钟预设和自定义时长。"]),
        ("第三阶段：优化安装交付", ["替换正式应用图标。", "优化 MSI 文件名和发布 ZIP 文件结构。", "同步更新使用手册中的安装和证书说明。"]),
    ]
    for phase, bullets in phases:
        doc.add_heading(phase, level=2)
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("五、验收标准", level=1)
    checks = [
        "首次启动时，用户能清楚看到隐私说明，并知道程序记录什么、不记录什么。",
        "结束学习后，用户可以在历史日报中重新查看之前的记录。",
        "AI 总结语气可选择，不同语气生成的内容风格有明显区别。",
        "番茄钟支持 25、40、50 分钟预设和自定义分钟数。",
        "安装包图标和文件名更正式，使用手册说明证书与 SmartScreen 限制。",
        "所有改动不改变隐私底线，不引入账号系统或云同步。",
    ]
    for check in checks:
        doc.add_paragraph(check, style="List Bullet")

    doc.add_heading("六、暂不纳入本轮的建议", level=1)
    doc.add_paragraph(
        "本轮暂不实现应用分类、白名单/黑名单和学习目标备注。这些功能有价值，但会影响专注度算法和数据结构设计，建议放入下一轮迭代。"
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("StudyPulse 项目改进材料")
    set_run_font(r, size=9)
    r.font.color.rgb = muted

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build_document()
