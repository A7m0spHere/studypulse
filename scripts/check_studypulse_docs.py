from pathlib import Path

from docx import Document


for path in sorted(Path("docs").glob("StudyPulse_*.docx")):
    doc = Document(path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    print(
        f"{path.name}: paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} chars={len(text)} first={doc.paragraphs[0].text if doc.paragraphs else ''}"
    )
